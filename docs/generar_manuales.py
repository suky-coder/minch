from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), '000000')
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_cover_page(doc, manual_type, version, date):
    for _ in range(4):
        doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MINCH SRL.')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Sistema de Información Contable')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SIC-MINCH')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(manual_type)
    run.bold = True
    run.font.size = Pt(18)
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Versi\xf3n: {version}')
    run.font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Fecha: {date}')
    run.font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('La Paz - Bolivia')
    run.font.size = Pt(13)
    doc.add_page_break()

def add_control_de_cambios(doc):
    doc.add_heading('Control de Cambios', level=1)
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    headers = ['Versi\xf3n', 'Fecha', 'Autor(es)', 'Descripci\xf3n del Cambio', 'Aprobador', 'Estado']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, 'D9D9D9')
    data = ['v1.0', 'Julio 2026', '[Nombre]', 'Creaci\xf3n inicial del documento', '[PO / QA Lead]', 'Borrador']
    for i, val in enumerate(data):
        cell = table.rows[1].cells[i]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9)

def add_indice(doc, items):
    doc.add_heading('\xcdndice General', level=1)
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

def add_glosario(doc, terms):
    doc.add_heading('Glosario de T\xe9rminos', level=1)
    table = doc.add_table(rows=len(terms) + 1, cols=2)
    table.style = 'Table Grid'
    for i, h in enumerate(['T\xe9rmino', 'Definici\xf3n']):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, 'D9D9D9')
    for r, (term, definition) in enumerate(terms):
        table.rows[r + 1].cells[0].text = term
        table.rows[r + 1].cells[1].text = definition
        for cell in table.rows[r + 1].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

def add_seccion_modulo(doc, num, titulo, descripcion, operaciones=None):
    doc.add_heading(f'{num}. {titulo}', level=2)
    p = doc.add_paragraph(descripcion)
    p.paragraph_format.space_after = Pt(4)
    if operaciones:
        p = doc.add_paragraph()
        run = p.add_run('Qu\xe9 puede hacer: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(operaciones)
        run.font.size = Pt(10)
    # Placeholder for screenshot
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add a bordered box for screenshot placeholder
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = '[Espacio para captura de pantalla]'
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(128, 128, 128)
            r.italic = True
        # Set background
        set_cell_shading(cell, 'F2F2F2')
    # Set cell height
    tr = table.rows[0]
    tr.height = Cm(6)
    doc.add_paragraph('')

def add_faq(doc):
    doc.add_heading('Soluci\xf3n de Problemas (FAQ)', level=1)
    faqs = [
        ('No puedo iniciar sesi\xf3n',
         'Verifique que su correo y contrase\xf1a sean correctos. Si sigue sin poder ingresar, solicite al administrador que restablezca su contrase\xf1a.'),
        ('No veo algunas opciones en el men\xfa',
         'Esto significa que su rol no tiene permisos para esas opciones. Solicite al administrador que revise sus permisos.'),
        ('Al guardar un movimiento aparece un error de saldo insuficiente',
         'Revise que el monto ingresado no supere el saldo disponible de la cuenta o caja seleccionada.'),
        ('No encuentro un proveedor o cliente en la lista',
         'Use el buscador escribiendo el nombre o la c\xe9dula de identidad. Si a\xfan no aparece, puede crear uno nuevo desde el formulario.'),
        ('El PDF no se descarga',
         'Verifique que su navegador no est\xe9 bloqueando las ventanas emergentes del sitio.'),
        ('Los datos no se guardan',
         'Revise que todos los campos obligatorios (marcados con *) est\xe9n llenos y que los valores sean correctos.'),
        ('No puedo eliminar un registro',
         'Algunos registros no se pueden eliminar si tienen movimientos asociados. Primero elimine los movimientos relacionados.'),
        ('El sistema se ve diferente o los botones no funcionan',
         'Intente actualizar la p\xe1gina con F5 o vac\xede la memoria cach\xe9 del navegador.'),
    ]
    for i, (problem, solution) in enumerate(faqs, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {problem}')
        run.bold = True
        run.font.size = Pt(10)
        p = doc.add_paragraph(f'   {solution}')
        p.paragraph_format.space_after = Pt(6)

def add_tabla_avance(doc):
    doc.add_heading('Avance en la Elaboraci\xf3n del Manual de Usuario', level=1)
    headers = ['Secci\xf3n / M\xf3dulo', 'Rol(es) Involucrados', 'Tareas Doc. (N\xb0)', 'Capturas (N\xb0)', 'Estado', 'Responsable']
    rows = [
        ['Introducci\xf3n y Perfiles', 'Todos', '4', '1', 'Revisado', 'Redactor T\xe9cnico'],
        ['Inicio de Sesi\xf3n', 'Todos', '3', '1', 'Borrador', 'Dev Frontend'],
        ['Dashboard', 'Todos', '2', '1', 'Borrador', 'Dev Frontend'],
        ['Usuarios', 'Administrador', '4', '2', 'Borrador', 'Dev Backend'],
        ['Roles', 'Administrador', '3', '1', 'Borrador', 'Dev Backend'],
        ['Permisos', 'Administrador', '3', '1', 'Borrador', 'Dev Backend'],
        ['Impuestos', 'Administrador, Contador', '4', '1', 'Borrador', 'Dev Backend'],
        ['Proveedores', 'Contador, Auxiliar', '3', '1', 'Borrador', 'Dev Backend'],
        ['Clientes', 'Contador, Auxiliar', '3', '1', 'Borrador', 'Dev Backend'],
        ['Cooperativas', 'Contador, Auxiliar', '4', '1', 'Borrador', 'Dev Backend'],
        ['Cuentas Bancarias', 'Administrador', '4', '1', 'Borrador', 'Dev Backend'],
        ['Libro de Bancos', 'Contador, Auxiliar', '5', '3', 'No Iniciado', 'Dev Fullstack'],
        ['Caja Chica', 'Auxiliar', '4', '2', 'No Iniciado', 'Dev Frontend'],
        ['Estados de Cuenta', 'Contador', '3', '2', 'No Iniciado', 'Dev Fullstack'],
        ['Retenciones', 'Contador', '6', '3', 'No Iniciado', 'Dev Backend'],
        ['Reportes', 'Gerente, Contador', '4', '2', 'No Iniciado', 'Dev Fullstack'],
        ['Perfil y Configuraci\xf3n', 'Todos', '3', '1', 'No Iniciado', 'Dev Frontend'],
        ['Soluci\xf3n de Problemas (FAQ)', 'Todos', '8', '0', 'Pendiente', 'Soporte'],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
        set_cell_shading(cell, 'D9D9D9')
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)

# ============================================================
# MANUAL DE USUARIO
# ============================================================
def generar_manual_usuario():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_cover_page(doc, 'Manual de Usuario', '1.0', 'Julio 2026')
    add_control_de_cambios(doc)
    doc.add_page_break()

    # Índice
    add_indice(doc, [
        '1. Introducci\xf3n y Perfiles de Usuario',
        '2. Inicio de Sesi\xf3n',
        '3. Dashboard (Panel Principal)',
        '4. Usuarios',
        '5. Roles',
        '6. Permisos',
        '7. Impuestos',
        '8. Proveedores',
        '9. Clientes',
        '10. Cooperativas',
        '11. Cuentas Bancarias',
        '12. Libro de Bancos',
        '13. Caja Chica',
        '14. Estados de Cuenta',
        '15. Retenciones',
        '16. Reportes',
        '17. Perfil y Configuraci\xf3n',
        '18. Soluci\xf3n de Problemas (FAQ)',
        '19. Avance de Elaboraci\xf3n',
    ])
    doc.add_page_break()

    # Glosario
    add_glosario(doc, [
        ('Dashboard', 'Pantalla principal con el resumen del sistema'),
        ('Retenci\xf3n', 'Descuento de impuestos que se aplica al pagar a un proveedor'),
        ('Caja Chica', 'Dinero en efectivo para gastos menores del d\xeda a d\xeda'),
        ('Libro de Bancos', 'Registro de todos los movimientos que entran y salen de las cuentas bancarias'),
        ('Movimiento D\xe9bito (D)', 'Ingreso de dinero a una cuenta o caja'),
        ('Movimiento Cr\xe9dito (C)', 'Salida de dinero de una cuenta o caja'),
        ('Rol', 'Cargo o puesto que tiene una persona dentro del sistema'),
        ('Permiso', 'Acci\xf3n espec\xedfica que un rol puede realizar (ver, crear, editar, eliminar)'),
        ('Proveedor', 'Empresa o persona que vende servicios o bienes a MINCH'),
        ('Cliente', 'Empresa o persona que compra productos de MINCH'),
        ('Cooperativa', 'Cooperativa minera que trabaja con MINCH'),
        ('Estado de Cuenta', 'Informe que muestra el saldo y los movimientos de un proveedor o cliente'),
    ])
    doc.add_page_break()

    # 1. Introducción
    doc.add_heading('1. Introducci\xf3n y Perfiles de Usuario', level=1)
    p = doc.add_paragraph('El sistema SIC-MINCH cuenta con diferentes niveles de acceso seg\xfan el cargo de cada persona. '
                          'A continuaci\xf3n se describen los perfiles de usuario:')
    profiles = [
        ('Auxiliar Contable (R1)', 'Puede registrar movimientos de caja chica y libro de bancos. Consulta saldos y movimientos.'),
        ('Contador General (R2)', 'Puede registrar retenciones, generar recibos PDF, consultar el libro de bancos y generar estados de cuenta.'),
        ('Administrador del Sistema (R3)', 'Tiene acceso completo. Puede crear usuarios, asignar roles y permisos, configurar impuestos y cuentas.'),
        ('Gerente General (R4)', 'Puede ver el dashboard, generar reportes con filtros y exportar a PDF o Excel.'),
    ]
    for role, desc in profiles:
        p = doc.add_paragraph()
        run = p.add_run(f'{role}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)
    doc.add_page_break()

    # Modules 2-17
    modules = [
        ('2. Inicio de Sesi\xf3n',
         'Para ingresar al sistema necesita un usuario y contrase\xf1a proporcionados por el administrador. '
         'Escriba su correo electr\xf3nico y contrase\xf1a en la pantalla de inicio y presione el bot\xf3n "Iniciar Sesi\xf3n". '
         'Si olvid\xf3 su contrase\xf1a, solicite al administrador que la restablezca.',
         'Todos los usuarios pueden iniciar sesi\xf3n.'),
        ('3. Dashboard (Panel Principal)',
         'Al ingresar al sistema ver\xe1 esta pantalla con un resumen de toda la informaci\xf3n importante. '
         'Aqu\xed puede ver la cantidad de proveedores y cooperativas registrados, los saldos de caja y bancos, '
         'los \xfaltimos movimientos y retenciones, y gr\xe1ficos que muestran la actividad del negocio.',
         'Todos los roles pueden ver esta pantalla. No se pueden modificar datos desde aqu\xed.'),
        ('4. Usuarios',
         'Este m\xf3dulo permite administrar las personas que pueden ingresar al sistema. '
         'Puede crear un nuevo usuario ingresando sus datos personales (nombre, apellido, c\xe9dula, correo, tel\xe9fono) '
         'y asignarle un rol. Tambi\xe9n puede editar los datos de un usuario existente o eliminarlo.',
         'Solo el Administrador puede crear, editar y eliminar usuarios.'),
        ('5. Roles',
         'Los roles son los cargos o puestos que tienen los usuarios (Administrador, Contador, Auxiliar, Gerente). '
         'Aqu\xed puede crear nuevos roles o cambiar el nombre de los existentes. '
         'Los roles permiten organizar qu\xe9 puede hacer cada persona en el sistema.',
         'Solo el Administrador puede crear, editar y eliminar roles.'),
        ('6. Permisos',
         'En este m\xf3dulo se define qu\xe9 puede hacer cada rol dentro del sistema. '
         'Seleccione un rol de la lista y luego marque o desmarque los permisos que desea asignarle. '
         'Por ejemplo, puede permitir que el rol "Auxiliar" solo vea informaci\xf3n pero no pueda eliminar datos. '
         'Use los botones "Asignar todo" o "Revocar todo" para cambios r\xe1pidos.',
         'Solo el Administrador puede asignar permisos.'),
        ('7. Impuestos',
         'Aqu\xed se registran los impuestos que se usan en las retenciones. '
         'El sistema ya incluye los impuestos RC-IVA (13%), IUE (5%) e IT (3%). '
         'Puede crear nuevos impuestos, editar su nombre, iniciales, porcentaje y tipo, o eliminarlos. '
         'Los tipos disponibles son: Servicios (S), Bienes (G) o Todos (A).',
         'El Administrador y Contador pueden gestionar impuestos.'),
        ('8. Proveedores',
         'Registra las empresas o personas que venden servicios o bienes a MINCH. '
         'Puede crear un nuevo proveedor ingresando su nombre completo, c\xe9dula de identidad y una descripci\xf3n. '
         'Tambi\xe9n puede editar sus datos o eliminarlo. '
         'Los proveedores se usan al registrar retenciones, movimientos bancarios y de caja.',
         'El Contador y Auxiliar pueden gestionar proveedores.'),
        ('9. Clientes',
         'Registra las empresas o personas que compran productos de MINCH. '
         'Adem\xe1s del nombre y c\xe9dula de identidad, puede adjuntar un archivo (como un contrato) '
         'y asociar el cliente a una cooperativa minera.',
         'El Contador y Auxiliar pueden gestionar clientes.'),
        ('10. Cooperativas',
         'Registra las cooperativas mineras con las que trabaja MINCH. '
         'Puede ingresar datos como el nombre, NIT, NIM, concesi\xf3n, mina (vocamina), municipio '
         'y los porcentajes de aporte a la cooperativa y a COMIBOL.',
         'El Contador y Auxiliar pueden gestionar cooperativas.'),
        ('11. Cuentas Bancarias',
         'Aqu\xed se registran las cuentas bancarias de MINCH. '
         'Puede crear una cuenta con su n\xfamero, nombre, tipo de moneda (Bolivianos, D\xf3lares o Euros), '
         'una sigla y un color para identificarla f\xe1cilmente en las listas.',
         'Solo el Administrador puede gestionar cuentas bancarias.'),
        ('12. Libro de Bancos',
         'Este m\xf3dulo registra todo el movimiento de dinero en las cuentas bancarias. '
         'Seleccione un mes y una cuenta para ver sus movimientos. '
         'Puede crear un nuevo movimiento eligiendo si es ingreso (D\xe9bito) o egreso (Cr\xe9dito), '
         'el m\xe9todo de pago (Transferencia o Cheque) y el proveedor. '
         'Tambi\xe9n puede editar o eliminar movimientos, descargar un comprobante PDF de cada movimiento '
         'o un estado de cuenta mensual en PDF y Excel.',
         'El Contador y Auxiliar pueden ver y gestionar movimientos. El Administrador configura las cuentas.'),
        ('13. Caja Chica',
         'Registra los gastos menores que se pagan en efectivo. '
         'Seleccione un mes para ver los movimientos del libro de caja. '
         'Puede crear un nuevo gasto (egreso) o ingreso, editarlo o eliminarlo. '
         'Tambi\xe9n puede descargar un reporte mensual en PDF.',
         'El Auxiliar Contable puede gestionar la caja chica.'),
        ('14. Estados de Cuenta',
         'Muestra cu\xe1nto se le debe a cada proveedor o cu\xe1nto debe cada cliente. '
         'Puede elegir entre la pesta\xf1a de "Proveedores" o "Clientes". '
         'Al hacer clic en el icono de ojo de una persona, puede ver el detalle de todos sus movimientos '
         'y el saldo actual. Puede agregar movimientos manuales y descargar un PDF con el estado de cuenta completo.',
         'El Contador puede gestionar y ver estados de cuenta.'),
        ('15. Retenciones',
         'Este m\xf3dulo permite crear recibos de retenci\xf3n de impuestos. '
         'Seleccione el tipo (Servicios o Bienes), busque el proveedor escribiendo su nombre o c\xe9dula, '
         'ingrese el monto pactado y el sistema calcular\xe1 autom\xe1ticamente los descuentos de cada impuesto. '
         'Puede ver la lista de retenciones por mes y tipo, descargar cada recibo en PDF '
         'y exportar la lista mensual a Excel.',
         'El Contador puede crear, editar y eliminar retenciones.'),
        ('16. Reportes',
         'Aqu\xed puede generar reportes de Retenciones, Caja Chica y Libro de Bancos. '
         'Seleccione las fechas de inicio y fin, aplique los filtros que necesite (tipo, cuenta) '
         'y presione "Consultar". Luego puede descargar el reporte en PDF o Excel.',
         'El Gerente y Contador pueden ver y exportar reportes.'),
        ('17. Perfil y Configuraci\xf3n',
         'En esta secci\xf3n puede cambiar sus datos personales como nombre, apellido, '
         'c\xe9dula de identidad, correo electr\xf3nico, tel\xe9fono y fecha de nacimiento. '
         'Tambi\xe9n puede cambiar su contrase\xf1a y configurar el tema claro u oscuro del sistema.',
         'Todos los usuarios pueden editar su propio perfil.'),
    ]
    for title, desc, ops in modules:
        add_seccion_modulo(doc, title.split('. ')[0], title.split('. ', 1)[1] if '. ' in title else title, desc, ops)
        doc.add_page_break()

    # FAQ
    add_faq(doc)
    doc.add_page_break()

    # Tabla de Avance
    add_tabla_avance(doc)

    path = os.path.join(OUTPUT_DIR, 'manual-usuario.docx')
    doc.save(path)
    print(f'[OK] {path}')

# ============================================================
# MANUAL DE INSTALACIÓN Y DESPLIEGUE
# ============================================================
def generar_manual_instalacion():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_cover_page(doc, 'Manual de Instalaci\xf3n y Despliegue', '1.0', 'Julio 2026')
    add_control_de_cambios(doc)
    doc.add_page_break()

    add_indice(doc, [
        '1. Requisitos Previos',
        '2. Instalaci\xf3n Paso a Paso',
        '3. Configuraci\xf3n de Seguridad',
        '4. Verificaci\xf3n del Despliegue',
        '5. Mantenimiento',
    ])
    doc.add_page_break()

    add_glosario(doc, [
        ('Servidor', 'Computadora donde se instala el sistema para que otros usuarios puedan acceder'),
        ('Base de Datos', 'Lugar donde se almacena toda la informaci\xf3n del sistema'),
        ('SSL', 'Certificado de seguridad para que la conexi\xf3n sea segura (https)'),
        ('Firewall', 'Barrera de seguridad que controla el tr\xe1fico de red'),
        ('Migraci\xf3n', 'Proceso que crea las tablas necesarias en la base de datos'),
        ('Seeder', 'Datos de ejemplo que se cargan para probar el sistema'),
        ('Repositorio', 'Lugar donde se almacena el c\xf3digo fuente del sistema (Git)'),
        ('Variable de Entorno', 'Configuraci\xf3n que el sistema necesita para funcionar (conexi\xf3n a BD, correo, etc.)'),
    ])
    doc.add_page_break()

    # 1. Requisitos Previos
    doc.add_heading('1. Requisitos Previos', level=1)
    doc.add_heading('Hardware', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    for i, h in enumerate(['Componente', 'M\xednimo']):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, 'D9D9D9')
    hw = [('Procesador', '2 n\xfacleos'), ('Memoria RAM', '4 GB'), ('Disco Duro', '20 GB libres')]
    for r, (comp, val) in enumerate(hw):
        table.rows[r + 1].cells[0].text = comp
        table.rows[r + 1].cells[1].text = val

    doc.add_heading('Software', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    for i, h in enumerate(['Componente', 'Versi\xf3n']):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, 'D9D9D9')
    sw = [('Sistema Operativo', 'Ubuntu 22.04 / Windows Server 2022'),
          ('PHP', '8.3'),
          ('MySQL', '8.0'),
          ('Composer', '2.x'),
          ('Node.js', '20.x'),
          ('Nginx', '\xdaltima estable')]
    for r, (comp, ver) in enumerate(sw):
        table.rows[r + 1].cells[0].text = comp
        table.rows[r + 1].cells[1].text = ver

    doc.add_heading('Puertos de Red', level=2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    for i, h in enumerate(['Puerto', 'Uso']):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, 'D9D9D9')
    ports = [('80/443', 'HTTP/HTTPS (navegador)'), ('3306', 'MySQL (solo interno)')]
    for r, (port, usage) in enumerate(ports):
        table.rows[r + 1].cells[0].text = port
        table.rows[r + 1].cells[1].text = usage
    doc.add_page_break()

    # 2. Instalación
    doc.add_heading('2. Instalaci\xf3n Paso a Paso', level=1)
    steps = [
        ('Descargar el c\xf3digo fuente',
         'Abra una terminal y ejecute:\n\n   git clone <url-del-repositorio>\n   cd minch'),
        ('Instalar dependencias de PHP',
         'Ejecute:\n\n   composer install'),
        ('Configurar archivo de entorno',
         'Copie el archivo de ejemplo y ed\xedtelos:\n\n   cp .env.example .env\n\nLuego edite el archivo .env con los datos de su base de datos (nombre, usuario, contrase\xf1a).'),
        ('Generar clave del sistema',
         'Ejecute:\n\n   php artisan key:generate'),
        ('Crear las tablas en la base de datos',
         'Ejecute:\n\n   php artisan migrate --seed\n\nEsto crear\xe1 todas las tablas y cargar\xe1 datos de ejemplo (impuestos, usuarios, cuentas bancarias).'),
        ('Instalar dependencias de JavaScript y compilar',
         'Ejecute:\n\n   npm install\n   npm run build'),
        ('Iniciar el sistema',
         'Para desarrollo ejecute:\n\n   php artisan serve\n\nPara producci\xf3n, configure Nginx para que apunte a la carpeta "public" del proyecto.'),
    ]
    for title, content in steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        p = doc.add_paragraph(content)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # 3. Seguridad
    doc.add_heading('3. Configuraci\xf3n de Seguridad', level=1)
    sec_steps = [
        ('Certificado SSL',
         'Para que los usuarios accedan de forma segura (https), instale un certificado SSL.\n\n'
         'Si tiene un dominio, use Let\'s Encrypt:\n\n   sudo apt install certbot python3-certbot-nginx\n   sudo certbot --nginx\n\n'
         'Si es solo para pruebas locales, puede usar un certificado autofirmado.'),
        ('Firewall',
         'Active el firewall para permitir solo los puertos necesarios:\n\n'
         '   sudo ufw allow 80/tcp\n   sudo ufw allow 443/tcp\n   sudo ufw enable'),
        ('Base de Datos',
         'Cree un usuario dedicado para la base de datos:\n\n'
         '   CREATE USER \'minch\'@\'localhost\' IDENTIFIED BY \'contrase\xf1a_segura\';\n'
         '   GRANT ALL PRIVILEGES ON minch_db.* TO \'minch\'@\'localhost\';\n'
         '   FLUSH PRIVILEGES;'),
    ]
    for title, content in sec_steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        p = doc.add_paragraph(content)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # 4. Verificación
    doc.add_heading('4. Verificaci\xf3n del Despliegue', level=1)
    p = doc.add_paragraph('Use la siguiente lista para verificar que la instalaci\xf3n fue exitosa:')

    headers = ['Paso / Secci\xf3n', 'Comandos / Acciones documentadas', '\xbfEs reproducible?', 'Estado', 'Observaciones']
    rows = [
        ['Prerrequisitos', 'PHP 8.3, MySQL 8.0, Composer 2.x, Node.js 20.x instalados', 'S\xed', 'OK', 'Verificar con php -v, mysql --version'],
        ['Variables de Entorno', 'Archivo .env con 20 variables configuradas', 'S\xed', 'OK', 'Revisar DB_HOST, DB_PORT, DB_DATABASE'],
        ['Dependencias PHP', 'composer install sin errores', 'S\xed', 'OK', 'Puede tardar 2-3 minutos'],
        ['Migraciones y Semillas', 'php artisan migrate --seed', 'S\xed', 'OK', 'Crea tablas + datos demo'],
        ['Compilaci\xf3n de Assets', 'npm run build', 'S\xed', 'OK', 'Tarda ~30 segundos'],
        ['Certificado SSL', 'certbot --nginx', 'S\xed (con dominio)', 'OK', 'En local usar autofirmado'],
        ['Prueba de Funcionamiento', 'Acceder a la URL principal', 'S\xed', 'OK', 'Debe aparecer la pantalla de login'],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
        set_cell_shading(cell, 'D9D9D9')
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Nota: ')
    run.bold = True
    run = p.add_run('Se recomienda realizar esta verificaci\xf3n despu\xe9s de cada actualizaci\xf3n del sistema.')

    doc.add_page_break()

    # 5. Mantenimiento
    doc.add_heading('5. Mantenimiento', level=1)
    maint = [
        ('Respaldos de Base de Datos',
         'Programe un respaldo autom\xe1tico diario usando cron:\n\n'
         '   0 2 * * * /usr/bin/mysqldump -u minch -p minch_db > /backups/minch_$(date +\\%Y\\%m\\%d).sql'),
        ('Limpiar Cach\xe9',
         'Si el sistema se vuelve lento o muestra informaci\xf3n desactualizada, limpie la cach\xe9:\n\n'
         '   php artisan cache:clear\n   php artisan view:clear\n   php artisan config:clear'),
        ('Actualizar el Sistema',
         'Para actualizar a una nueva versi\xf3n:\n\n'
         '   git pull\n   composer install\n   php artisan migrate\n   npm run build'),
    ]
    for title, content in maint:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        p = doc.add_paragraph(content)
        p.paragraph_format.space_after = Pt(6)

    path = os.path.join(OUTPUT_DIR, 'manual-instalacion.docx')
    doc.save(path)
    print(f'[OK] {path}')

# ============================================================
# MANUAL DE BASE DE DATOS
# ============================================================
def generar_manual_bd():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_cover_page(doc, 'Manual de Base de Datos', '1.0', 'Julio 2026')
    add_control_de_cambios(doc)
    doc.add_page_break()

    add_indice(doc, [
        '1. Diagrama Entidad-Relaci\xf3n (DER)',
        '2. L\xf3gica de Negocio en Base de Datos',
        '3. Plan de Respaldo y Recuperaci\xf3n',
        '4. \xcdndices y Optimizaci\xf3n',
        '5. Documentaci\xf3n de Entidades y Procedimientos',
    ])
    doc.add_page_break()

    add_glosario(doc, [
        ('DER', 'Diagrama que muestra las tablas de la base de datos y c\xf3mo se relacionan entre s\xed'),
        ('FK (Clave For\xe1nea)', 'Campo que conecta una tabla con otra'),
        ('PK (Clave Primaria)', 'Identificador \xfanico de cada registro en una tabla'),
        ('Migraci\xf3n', 'Archivo de Laravel que define c\xf3mo crear o modificar una tabla'),
        ('lockForUpdate()', 'Mecanismo que evita que dos personas creen el mismo n\xfamero de recibo al mismo tiempo'),
        ('mysqldump', 'Comando para hacer una copia de seguridad de la base de datos'),
        ('\xcdndice', 'Estructura que acelera las b\xfasquedas en la base de datos'),
        ('Funci\xf3n de Ventana', 'Forma de calcular el saldo acumulado en cada movimiento sin usar tablas temporales'),
    ])
    doc.add_page_break()

    # 1. DER
    doc.add_heading('1. Diagrama Entidad-Relaci\xf3n (DER)', level=1)
    p = doc.add_paragraph('A continuaci\xf3n se muestra el diagrama de la base de datos del sistema SIC-MINCH:')
    doc.add_paragraph('')
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = '[Espacio para insertar el Diagrama Entidad-Relaci\xf3n (DER)]'
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(128, 128, 128)
            r.italic = True
    set_cell_shading(cell, 'F2F2F2')
    table.rows[0].height = Cm(10)
    doc.add_paragraph('')
    p = doc.add_paragraph('Nota: El DER puede generarse con herramientas como MySQL Workbench o DBeaver a partir del archivo de migraciones.')
    doc.add_page_break()

    # 2. Lógica de Negocio
    doc.add_heading('2. L\xf3gica de Negocio en Base de Datos', level=1)

    doc.add_heading('2.1 Numeraci\xf3n Autom\xe1tica de Recibos', level=2)
    p = doc.add_paragraph('El sistema asigna n\xfameros de forma autom\xe1tica a los recibos usando un mecanismo '
                          'de bloqueo (lockForUpdate) para evitar que dos usuarios obtengan el mismo n\xfamero al mismo tiempo:')
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(['Documento', 'Formato del N\xfamero', 'Periodo']):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, 'D9D9D9')
    data = [
        ['Retenciones', 'MES-001 (ej. ABR-001)', 'Mes calendario'],
        ['Transacciones Bancarias', 'N\xfamero correlativo por a\xf1o fiscal', 'A\xf1o fiscal (Oct-Sep)'],
        ['Caja Chica', 'N\xfamero correlativo por mes', 'Mes calendario'],
    ]
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val
            for p in table.rows[r + 1].cells[c].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_heading('2.2 C\xe1lculo de Retenciones', level=2)
    p = doc.add_paragraph('Cuando se crea una retenci\xf3n, el sistema calcula autom\xe1ticamente los descuentos:\n\n'
                          '   Total Bruto = Monto / (1 - Suma de tasas / 100)\n'
                          '   Descuento = Total Bruto x Tasa / 100\n\n'
                          'Seg\xfan el tipo de retenci\xf3n:\n'
                          '   - Servicios (S): RC-IVA (13%) + IT (3%) = 16%\n'
                          '   - Bienes (G): IUE (5%) + IT (3%) = 8%')

    doc.add_heading('2.3 C\xe1lculo de Saldos en Tiempo Real', level=2)
    p = doc.add_paragraph('Los saldos del libro de bancos y caja chica se calculan autom\xe1ticamente '
                          'usando funciones de ventana de MySQL. Al crear, editar o eliminar un movimiento, '
                          'el sistema recalcula el saldo desde la fecha del movimiento hasta el final.\n\n'
                          '   Saldo = SUM(CASE WHEN tipo IN ("D","B") THEN monto ELSE -monto END)\n'
                          '            OVER (ORDER BY fecha, id)')

    doc.add_page_break()

    # 3. Respaldo y Recuperación
    doc.add_heading('3. Plan de Respaldo y Recuperaci\xf3n (BDR)', level=1)

    doc.add_heading('3.1 Respaldo Completo', level=2)
    p = doc.add_paragraph('Para hacer una copia de seguridad completa de la base de datos, use el comando mysqldump:\n\n'
                          '   mysqldump -u usuario -p contrase\xf1a minch_db > respaldo_$(date +%Y%m%d).sql\n\n'
                          'Este comando crea un archivo .sql con todas las tablas y datos.')

    doc.add_heading('3.2 Respaldo Autom\xe1tico (Cron)', level=2)
    p = doc.add_paragraph('Programe un respaldo diario autom\xe1tico agregando esta l\xednea al crontab:\n\n'
                          '   0 3 * * * /usr/bin/mysqldump -u minch -p contrase\xf1a minch_db | gzip > /backups/minch_$(date +\\%Y\\%m\\%d).sql.gz\n\n'
                          'Se recomienda:\n'
                          '   - Conservar los respaldos de los \xfaltimos 7 d\xedas\n'
                          '   - Almacenar los respaldos en un disco diferente al de la base de datos\n'
                          '   - Realizar una copia externa (nube o NAS) una vez por semana')

    doc.add_heading('3.3 Restauraci\xf3n de Base de Datos', level=2)
    p = doc.add_paragraph('Para restaurar la base de datos desde un respaldo:\n\n'
                          '   mysql -u usuario -p minch_db < respaldo_20260701.sql\n\n'
                          'Pasos recomendados:\n'
                          '   1. Crear una base de datos vac\xeda: CREATE DATABASE minch_db;\n'
                          '   2. Restaurar el respaldo: mysql -u minch -p minch_db < respaldo.sql\n'
                          '   3. Verificar que las tablas se crearon: mysql -u minch -p -e "USE minch_db; SHOW TABLES;"')

    doc.add_page_break()

    # 4. Índices
    doc.add_heading('4. \xcdndices y Optimizaci\xf3n', level=1)
    p = doc.add_paragraph('Los siguientes \xedndices se crearon para mejorar la velocidad de las consultas m\xe1s frecuentes:')
    headers = ['\xcdndice', 'Tabla', 'Campo(s)', 'Prop\xf3sito']
    rows = [
        ['idx_movements_date_type', 'movements', 'date, type', 'Acelera la b\xfasqueda de movimientos por fecha y tipo'],
        ['idx_retentions_code', 'retentions', 'code', 'Garantiza que el c\xf3digo de retenci\xf3n sea \xfanico y acelera la b\xfasqueda'],
        ['idx_transactions_account_date', 'transactions', 'account_id, date', 'Acelera la consulta de movimientos por cuenta y fecha'],
        ['idx_boxes_date', 'boxes', 'date', 'Acelera la b\xfasqueda de movimientos de caja por fecha'],
        ['idx_movements_person_id', 'movements', 'person_id', 'Acelera la b\xfasqueda de movimientos de un proveedor/cliente'],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, 'D9D9D9')
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_page_break()

    # 5. Tabla de Documentación
    doc.add_heading('5. Documentaci\xf3n de Entidades y Procedimientos', level=1)
    headers = ['Componente', 'Nombre / Identificador', 'Descripci\xf3n', 'Prop\xf3sito / Estado', 'Doc. Adjunta']
    rows = [
        ['Tabla Principal', 'retentions',
         'Recibos de retenci\xf3n: tipo (S/G), monto, c\xf3digo correlativo, proveedor, usuario. Incluye resumen de descuentos en JSON.',
         'Tabla cr\xedtica', 'S\xed'],
        ['Tabla Principal', 'transactions',
         'Movimientos bancarios: cuenta, tipo de pago (T/CH), n\xfamero de cheque, n\xfamero fiscal por a\xf1o.',
         'Tabla cr\xedtica', 'S\xed'],
        ['Tabla Principal', 'boxes',
         'Movimientos de caja chica: tipo (D/C), monto, fecha, n\xfamero correlativo mensual.',
         'Tabla cr\xedtica', 'S\xed'],
        ['Tabla Principal', 'movements',
         'Tabla que unifica todos los movimientos (D/C/B). Se conecta con transactions y boxes.',
         'Tabla cr\xedtica', 'S\xed'],
        ['Tabla Maestra', 'taxes',
         'Cat\xe1logo de impuestos: nombre, iniciales (RC-IVA, IUE, IT), porcentaje, tipo (S/G/A).',
         'Tabla maestra', 'S\xed'],
        ['Tabla Maestra', 'accounts',
         'Cuentas bancarias: nombre, n\xfamero, iniciales, moneda (BOB/USD/EUR), color.',
         'Tabla maestra', 'S\xed'],
        ['Tabla Maestra', 'persons / suppliers / customers',
         'Datos de personas. Se conectan uno a uno con proveedores y clientes.',
         'Tabla maestra', 'S\xed'],
        ['Evento (Modelo)', 'Retention::creating()',
         'Asigna c\xf3digo secuencial por mes usando lockForUpdate().',
         'L\xf3gica de negocio clave', 'S\xed'],
        ['Evento (Modelo)', 'Transaction::creating()',
         'Asigna n\xfamero fiscal secuencial por a\xf1o (octubre-septiembre) con lockForUpdate().',
         'L\xf3gica de negocio clave', 'S\xed'],
        ['Evento (Modelo)', 'Box::creating()',
         'Asigna n\xfamero de caja secuencial por mes calendario con lockForUpdate().',
         'L\xf3gica de negocio clave', 'S\xed'],
        ['Vista (SQL)', 'Saldo corriente con ventana',
         'Calcula saldo acumulado: SUM(amount) OVER (PARTITION BY account_id ORDER BY date, id)',
         'Optimizaci\xf3n', 'S\xed'],
        ['\xcdndice', 'idx_movements_date_type',
         '\xcdndice compuesto en movements por fecha y tipo.',
         'Optimizaci\xf3n', 'S\xed'],
        ['\xcdndice', 'idx_retentions_code',
         '\xcdndice \xfanico en retentions por c\xf3digo.',
         'Integridad', 'S\xed'],
        ['Script', 'backup_automatico.sh',
         'Script de shell para mysqldump cifrado con rotaci\xf3n de 7 d\xedas.',
         'Respaldo', 'S\xed'],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
        set_cell_shading(cell, 'D9D9D9')
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    path = os.path.join(OUTPUT_DIR, 'manual-bd.docx')
    doc.save(path)
    print(f'[OK] {path}')

# ============================================================
# MAIN
# ============================================================
if __name__ == '__main__':
    generar_manual_usuario()
    generar_manual_instalacion()
    generar_manual_bd()
    print('\nTodos los manuales fueron generados correctamente.')
